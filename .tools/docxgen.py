#!/usr/bin/env python3
"""
docxgen.py — write the six assignment descriptions as .docx files.

Requires python-docx:   pip install python-docx

Each project description is data below, so the six documents share one layout
and one voice.  They deliberately go further than the syllabus text in three
places, because the syllabus text is thin in exactly those places:

  * OPTIONAL FEATURES.  The syllabus mentions "optional switch" once and says
    nothing about for-loops, do-while, structs or strings, even though the
    course code supports several of them.  Each description now names the
    optional extensions that make sense at that milestone and says what
    claiming one requires.

  * CLARIFIED EXPECTATIONS.  Anywhere the original wording could be read two
    ways ("complete each phase", "extensive comments", "demonstrate the
    compiler working"), the description now says what specifically is being
    asked for.

  * DOCUMENTATION AND VIDEO.  Both are now stated as concrete requirements
    with a checklist, rather than a sentence.

No grading rubric is included; the instructor maintains those separately.
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    sys.exit("docxgen: needs python-docx.  Run:  pip install python-docx")

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
DOCS = os.path.join(ROOT, "docs")

ACCENT = RGBColor(0x7A, 0x4B, 0x1E)
MUTED = RGBColor(0x60, 0x5A, 0x52)

# ===========================================================================
# Shared blocks
# ===========================================================================

VIDEO_INTRO = (
    "Every team member records their OWN video. Redundancy across team members is "
    "expected and fine — each video must stand on its own, because each is watched "
    "on its own. Use Loom, Zoom, YouTube (unlisted), or any host that produces a "
    "link that works without a login. Submit the individual links.")

VIDEO_COMMON = [
    ("Show it working, first",
     "Open with the compiler running on a non-trivial program, end to end, including "
     "executing the generated assembly. Do this before any explanation. Everything "
     "you say afterwards lands differently once the viewer has seen the thing work."),
    ("Explain how the code works",
     "Walk through the code that does the work, on screen. Not line-by-line narration "
     "— the viewer can read. Explain the STRUCTURE: what each function is responsible "
     "for, what it hands to the next phase, and what would break if it did that "
     "differently."),
    ("Explain WHY it works that way",
     "Name at least one decision you made, say what the alternative was, and say why "
     "you chose as you did. A decision with no alternative was not a decision."),
    ("Outline your own contribution",
     "State plainly which parts you personally wrote or debugged. If you paired on "
     "something, say so."),
    ("Demonstrate correct execution",
     "Run the test programs. Show the output. If something does not work, show that "
     "too and say what you know about why — a known, diagnosed failure is worth far "
     "more than a hidden one."),
]

DOC_REQUIREMENTS = [
    "Every source file opens with a header comment saying which compiler PHASE it "
    "implements, what it receives from the previous phase, and what it produces for "
    "the next one.",
    "Every function that is not obvious from its name has a comment saying what it "
    "does and what its return value means. 'What it means' matters more than 'what it "
    "does' — for example, that generateTACExpr returns the NAME of the location "
    "holding the result.",
    "Every non-obvious decision carries a comment explaining WHY, not what. A comment "
    "that restates the code is worse than no comment; a comment explaining why the "
    "registers are flushed at a label is worth ten lines of narration.",
    "Every grammar rule with a non-trivial action is commented with the construct it "
    "recognises and the AST node it builds.",
    "Known limitations are recorded in comments where they bite, not only in the "
    "README.",
    "Commented-out code is deleted before submission. Version control is what history "
    "is for.",
]

GITHUB = [
    "Push all source to your PRIVATE GitHub repository and add the instructor as a "
    "collaborator.",
    "Commit as you go, not once at the end. The commit history is evidence of how the "
    "work was divided, and a single commit named 'final' is evidence of the opposite.",
    "Do NOT commit build artefacts: no .o files, no lex.yy.c, no parser.tab.c/.h, no "
    "compiled binaries, no .s or .tac output. A .gitignore is provided.",
    "Include the test programs you wrote, not only the ones supplied.",
]


# ===========================================================================
# The six projects
# ===========================================================================

PROJECTS = [
 dict(
  n=1, topic=1,
  slug="topic-1-lexical-analysis",
  name="Project 1 — Lexical Analyzer",
  window="Topic 1 · Sep 8 – Sep 13, 2026",
  due="Friday, September 13, 2026, 11:59 PM",
  clc=True,
  summary=(
    "Build the lexical analyzer for the language you will compile for the rest of the "
    "semester. This is the first of six milestones; every later project extends this "
    "same code base, so the decisions you make this week — particularly the token "
    "names — will be with you until December."),
  goal=(
    "A program that reads a source file, prints every token with its kind, its text "
    "and its exact location, reports every lexical error it finds, and certifies the "
    "file when there are none."),
  new=[
    ("Environment", "Unix/Linux with flex, bison and gcc installed and verified."),
    ("The token set", "Every keyword, operator and delimiter in the full course "
     "language — not only the ones Project 2 will use."),
    ("Location tracking", "Line AND column for every token and every error."),
    ("Error recovery", "Report all lexical errors in one run; do not stop at the first."),
  ],
  required=[
    "Evidence of access to Unix/Linux, as a screenshot with a caption explaining what "
    "it shows.",
    "Evidence that flex and bison are installed and working: the version output of "
    "each, plus a screenshot of a successful build. Captions required — a screenshot "
    "with no explanation demonstrates nothing.",
    "The scanner recognises EVERY keyword, operator, delimiter, identifier and integer "
    "literal in the provided grammar. Include the keywords used only in later topics "
    "(if, else, while, for, switch, case, default, break, return): adding them now "
    "costs one line each and saves editing this file in week 12.",
    "Each token is returned with the correct KIND and, where it matters, its TEXT. "
    "Identifiers and numbers carry their text; punctuation does not need to.",
    "Both comment forms are handled: // to end of line, and /* ... */ spanning lines. "
    "Comments produce no tokens but MUST still advance the line counter.",
    "An unterminated block comment is reported as an error, not silently swallowed.",
    "Lexical errors report the offending character with its exact LINE and COLUMN. "
    "Flex tracks lines for you; columns you must track yourself.",
    "After an error the scanner CONTINUES. A file with four bad characters produces "
    "four messages from one run.",
    "When there are no errors, the program says so explicitly and exits with status 0. "
    "On error it exits non-zero.",
    "A test suite of at least six input files, including at least two that are "
    "supposed to fail. Each test states its expected result in a header comment.",
  ],
  optional=[
    ("String literals", "Add \"...\" to the language. Decide what an unterminated "
     "string at end of line should do, and say why."),
    ("Character literals", "'a' and the escape sequences. Note what this forces the "
     "parser to deal with later."),
    ("Floating-point literals", "Recognising them is easy; note in your write-up what "
     "it would cost the rest of the compiler to support them."),
    ("Hex and octal integers", "0x1F, 0755. Where does the conversion happen — in the "
     "scanner action or later?"),
    ("A token-frequency report", "Print a count per token kind at the end. Cheap, and "
     "genuinely useful when you are debugging a grammar in Topic 2."),
  ],
  clarified=[
    ("\"Recognises every valid keyword, expression, or construct\"",
     "A scanner recognises TOKENS, not expressions or constructs — those are the "
     "parser's job in Topic 2. What is being asked for here is complete token "
     "coverage. Being able to state that boundary crisply is itself part of the "
     "assessment."),
    ("\"Returns the correct token and its kind\"",
     "Two things: the KIND (an enum value the parser will switch on) and the TEXT for "
     "the kinds where text carries meaning. Your token listing must show both."),
    ("\"Generates lexical errors ... including their exact location\"",
     "Line AND column. A line number alone is not an exact location on a line "
     "containing three statements."),
    ("\"Certifies that the source code is correct\"",
     "An explicit statement that the file is lexically clean, plus exit status 0. Be "
     "careful to make this claim only about VOCABULARY: a lexically clean file may "
     "still be nonsense, and your message should not imply otherwise."),
    ("Resubmission from CST-301",
     "You may reuse your CST-301 scanner as a starting point, but it must be extended "
     "to the full token set above and to the error and location requirements. A "
     "verbatim resubmission will not meet these requirements."),
  ],
  video=[
    ("Explain all six phases of compiler design",
     "Name each phase, what it takes in, what it produces, and the question it "
     "answers. Then say which of them you have built this week and which you have "
     "not. Two to three minutes, and it should be fluent — this is the framing for "
     "the whole semester."),
    ("Describe theoretical principles you put into practice",
     "Regular expressions and finite automata; longest-match and first-match rule "
     "resolution; the separation of vocabulary from grammar. Do not recite "
     "definitions — point at the place in YOUR scanner where each one shows up."),
    ("Demonstrate the working lexer",
     "Run it on a clean file and on a file with several errors. Show the token table "
     "and the error messages, and point at a column number to show it is correct."),
  ],
  start=[
    "Starter code: student/topic-1-lexical-analysis/lexer/",
    "Lecture notes: docs/topic-1-lexical-analysis/lecture-notes.html",
    "Build with make; self-test with make test.",
    "The TODOs in scanner.l are numbered in the order you should do them.",
  ]),

 dict(
  n=2, topic=2,
  slug="topic-2-minimal-compiler",
  name="Project 2 — Minimalist Language Compiler",
  window="Topic 2 · Sep 14 – Oct 11, 2026",
  due="Saturday, October 11, 2026, 11:59 PM",
  clc=True,
  summary=(
    "Build a complete compiler — all six phases — for a very small language. This is "
    "the largest project in the course by hours, and not because the language is "
    "hard. It is because everything you build here is the scaffolding that Projects 3 "
    "through 6 extend. Every later milestone is a new case in a switch statement you "
    "write this month."),
  goal=(
    "A compiler that takes a source file in the starter language and emits MIPS "
    "assembly that runs correctly in SPIM, with useful diagnostics for programs that "
    "are wrong."),
  new=[
    ("Phase 2 — parser", "A bison grammar that recognises the starter language and "
     "builds an abstract syntax tree."),
    ("Phase 3 — semantic analysis", "A symbol table, and checks for undeclared and "
     "duplicately declared variables."),
    ("Phase 4 — intermediate code", "Three-address code generated from the AST."),
    ("Phase 5 — optimization", "Constant folding and constant propagation, run to a "
     "fixed point."),
    ("Phase 6 — code generation", "MIPS assembly that SPIM will run."),
  ],
  required=[
    "The grammar accepts the starter language: integer declarations, assignment, "
    "addition, and print. Every rule builds an AST node.",
    "printAST produces a readable, correctly indented tree. This is not decoration — "
    "it is the instrument you will debug every later phase with, so make it good now.",
    "The symbol table records every declared variable and its storage location, and "
    "can be printed on demand.",
    "Semantic analysis rejects use of an undeclared variable and duplicate "
    "declaration in the same scope, each with a line number and the offending name.",
    "Three-address code is generated for every construct, written to a .tac file, and "
    "readable enough that you can hand-execute it.",
    "At least two optimizations are implemented and REPORTED: the compiler must say "
    "what it changed and how many instructions it removed.",
    "MIPS output assembles and runs in SPIM or QtSPIM and produces correct results.",
    "Syntax errors report the line and say what was expected — not merely 'syntax "
    "error'. At minimum, handle a missing semicolon after each statement form.",
    "A test suite of at least eight programs, including at least three that are "
    "supposed to fail (one lexical, one syntax, one semantic). Each states its "
    "expected result in a header comment.",
    "The compiler exits non-zero when compilation fails.",
  ],
  optional=[
    ("Subtraction, multiplication, division",
     "The starter grammar has only +. Adding the rest is four grammar rules and a "
     "precedence table — but do it only after the whole pipeline works end to end for "
     "+ alone. Note that Project 3 requires this anyway, so doing it early is credit "
     "toward the next milestone rather than extra work."),
    ("Parentheses in expressions",
     "One rule, no AST node. Worth doing for the discussion of why it needs no node."),
    ("Unary minus",
     "Requires %prec UMINUS. A good, small encounter with precedence declarations."),
    ("A --dot flag that emits GraphViz",
     "Write the AST as a .dot file and render it with GraphViz. Three lines of "
     "output code, and it produces the tree diagrams your video and your Lab Question "
     "6 answer both need."),
    ("Dead code elimination",
     "Beyond the two required optimizations. Note carefully which assignments are safe "
     "to remove and which are not."),
    ("Richer error recovery",
     "Error productions for more than the missing semicolon: unbalanced parentheses, "
     "a missing identifier after int, an assignment with no expression."),
  ],
  clarified=[
    ("\"Complete each phase and functionality outlined within the checklist\"",
     "All six phases must be implemented and demonstrable INDIVIDUALLY. Your compiler "
     "must be able to show its output at each stage — token stream, AST, symbol "
     "table, unoptimized TAC, optimized TAC, MIPS — because a compiler you cannot "
     "inspect phase by phase is a compiler you cannot debug."),
    ("\"For the minimal grammar provided\"",
     "The grammar is in docs/C-Minus-Grammar-Reference.md, Iteration 1, and in the "
     "lecture notes. You may extend it (see optional features) but the required "
     "grammar must work first."),
    ("\"Extensive comments in the code\"",
     "See the Code Documentation section below. 'Extensive' means the reader can "
     "follow the design, not that every line has a comment."),
    ("How much of this is one person's work",
     "This is a CLC assignment, but every team member's video must show them "
     "explaining code and answering for decisions. Divide by PHASE rather than by "
     "file, so that each person owns a boundary and has something coherent to "
     "explain."),
    ("What 'working' means for phase 6",
     "The generated .s file loads in SPIM without errors and prints the correct "
     "numbers. Assembly that looks plausible but does not run does not count."),
  ],
  video=[
    ("Describe leveraging the AST to generate IR and assembly",
     "Take one statement and follow it: AST node, then TAC, then MIPS. Show all three "
     "on screen for the same statement. This is the specific thing the assignment "
     "asks you to describe, and showing it is far more convincing than describing it."),
    ("Explain your optimizations and what they achieved",
     "Show the unoptimized and optimized .tac side by side, and quote the instruction "
     "counts."),
  ],
  start=[
    "Starter code: student/topic-2-minimal-compiler/compiler/",
    "Lecture notes: docs/topic-2-minimal-compiler/lecture-notes.html",
    "The scanner, all headers, the driver and the register allocator are given. The "
    "TODOs are the six phases.",
    "Suggested pace: week 2 grammar + AST, week 3 semantics, week 4 TAC, week 5 "
    "optimizer + MIPS. Get one statement all the way to running MIPS in week 2 — do "
    "not leave code generation until the last week.",
  ]),

 dict(
  n=3, topic=3,
  slug="topic-3-arrays-and-functions",
  name="Project 3 — Complex Variables and Functions",
  window="Topic 3 · Oct 12 – Nov 1, 2026",
  due="Saturday, November 1, 2026, 11:59 PM",
  clc=True,
  summary=(
    "Extend the Project 2 compiler with arrays, full arithmetic, functions and scope. "
    "Three of those four are additive — new cases in code you already wrote. The "
    "fourth, functions, introduces the activation record, which is the one genuinely "
    "new idea in this milestone and the thing every later project depends on."),
  goal=(
    "A compiler that handles arrays, arithmetic expressions with correct precedence, "
    "and functions with parameters, return values and recursion."),
  new=[
    ("Arrays", "Declaration with a constant size, element read and write, and arrays "
     "as parameters."),
    ("Full arithmetic", "+ - * / with correct precedence and associativity, "
     "parentheses, and unary minus."),
    ("Functions", "Definition, parameters, calls, return values, and recursion."),
    ("Scope", "A scope stack; parameters and locals shadow globals."),
    ("Activation records", "A real stack frame per call, with a saved return address."),
  ],
  required=[
    "Array declaration with a constant size, and element access on both sides of an "
    "assignment: arr[i] = arr[j] + 1.",
    "Arrays passed to functions BY REFERENCE. The callee receives a base address, not "
    "a copy. Your write-up must state this and explain why the parameter has no size.",
    "Global arrays and local arrays both work. They take different paths to their base "
    "address; your code generator must handle both and your video should show that you "
    "know which is which.",
    "All four arithmetic operators with correct precedence, plus parentheses and "
    "unary minus. 2 + 3 * 4 must be 14, and 17 - 5 - 3 must be 9.",
    "Function definitions with zero or more parameters, calls with matching arguments, "
    "and return values.",
    "Recursion works. factorial(5) is the minimum bar; if it returns anything other "
    "than 120 you have a calling-convention bug.",
    "Nested calls work: f(1, g(2)) and f(g(1), 2) must both be correct.",
    "A real activation record per call: the frame is pushed on entry and popped on "
    "exit, and $ra is saved in it. A callee's locals must not disturb its caller's.",
    "return jumps to the function epilogue. It must not fall through into the code "
    "after it.",
    "Semantic analysis catches: undeclared arrays, calls to undeclared functions, "
    "wrong argument counts, duplicate declarations in the same scope, and an array "
    "declared with a size of zero or less.",
    "Two-pass semantic analysis, so a function may be called before it is defined. If "
    "you choose the one-pass alternative, say so explicitly and document that forward "
    "references are not supported.",
    "All Project 2 tests still pass. A milestone that breaks the previous one is not a "
    "milestone.",
  ],
  optional=[
    ("Array bounds checking",
     "Generate a run-time check on every subscript. This requires the length to be "
     "available at run time — decide how (pass it, or store it before element 0) and "
     "explain what it costs in both code size and speed."),
    ("Multi-dimensional arrays",
     "int m[3][4]. The address arithmetic becomes base + (i*cols + j)*4. Worth doing "
     "if only to discover that the compiler now needs to remember the row length."),
    ("More than four arguments",
     "Pass the fifth and beyond on the stack. Decide who pops them — caller or callee "
     "— and say why. This is a real calling-convention design decision."),
    ("void functions",
     "A function with no return value. Requires a type on the function, which is the "
     "first step toward an actual type system."),
    ("Structures",
     "struct Point { int x; int y; }. A substantial extension: field offsets in the "
     "symbol table, a dotted-access AST node, and address arithmetic that composes "
     "with array indexing. Attempt only after everything required works, and document "
     "it thoroughly — this is the single most impressive optional feature at this "
     "milestone."),
    ("Global initialisers",
     "int total = 0; at file scope. Where does the initialisation happen — in .data, "
     "or as code before main?"),
  ],
  clarified=[
    ("\"Support for simple mathematical expressions\"",
     "All four arithmetic operators, with correct precedence and associativity, "
     "parentheses, and unary minus. Precedence must be demonstrably correct: your test "
     "suite must contain expressions whose value differs under the wrong precedence."),
    ("\"Support for arrays\"",
     "Declaration, element read, element write, and passing an array to a function. "
     "Bounds checking is optional (see above) but if you do not implement it, say so "
     "in your README as a known limitation."),
    ("\"Support for declaring and calling functions\"",
     "Including recursion and nested calls. A compiler that handles f(1) but not "
     "f(g(1)) has not met this requirement, and the difference between them is "
     "entirely in your calling convention."),
    ("\"Implement scope management methods\"",
     "Note carefully that there are TWO symbol tables in a compiler and they answer "
     "different questions: one for VISIBILITY during semantic analysis, one for "
     "STORAGE during code generation. Your write-up and your video should show you "
     "know the difference."),
    ("\"Integrate increasingly complex compiler components\"",
     "In practice this means: the Topic 2 tests must still pass. Run them. A "
     "regression here is a far more serious problem than a missing optional feature."),
  ],
  video=[
    ("Describe your approach to integrating increasingly complex components",
     "The specific question this project's video asks. A good answer names the "
     "interfaces that did NOT change — the AST node structure, the TAC instruction "
     "set, the phase boundaries — and explains why that was possible."),
    ("Show the activation record",
     "Compile a function, show the symbol table the compiler prints, and point at the "
     "prologue and epilogue in the generated assembly. Then run a recursive function "
     "and explain what is on the stack at maximum depth."),
  ],
  start=[
    "Starter code: student/topic-3-arrays-and-functions/compiler/",
    "It is your working Topic 2 compiler. It builds and passes the Topic 2 tests "
    "before you touch it.",
    "tests/ contains the five programs it must pass when you are done.",
    "The activation record is the hard part. Do the arithmetic and arrays first, then "
    "functions, and expect functions to take longer than the other three combined.",
  ]),

 dict(
  n=4, topic=4,
  slug="topic-4-loops",
  name="Project 4 — Loops and Optimization",
  window="Topic 4 · Nov 2 – Nov 22, 2026",
  due="Saturday, November 22, 2026, 11:59 PM",
  clc=True,
  summary=(
    "Add loops, and make the optimizer produce numbers you can defend. Loops are what "
    "make optimization matter: an instruction inside a loop that runs a thousand times "
    "costs a thousand instructions, so this is the first milestone where phase 5 stops "
    "being a formality."),
  goal=(
    "A compiler that lowers structured loops into labels and jumps, and an optimizer "
    "whose effect you have measured and can quantify."),
  new=[
    ("Relational operators", "< > <= >= == != , producing 1 or 0."),
    ("Loops", "At least one loop form, lowered to labels and conditional jumps."),
    ("break", "Leave the innermost loop early, with a semantic check that rejects it "
     "outside a loop."),
    ("A real optimizer", "Several techniques, run to a fixed point, reporting what "
     "each achieved."),
    ("Measurement", "Code size and executed-instruction counts, before and after."),
  ],
  required=[
    "At least ONE loop form fully implemented through all six phases: while, for, or "
    "repeat. Implementing more than one is optional (see below).",
    "The loop tests at the TOP: a loop whose condition is initially false executes "
    "zero times. Your test suite must contain such a program.",
    "Nested loops work, with distinct labels. A label generator that can repeat a name "
    "will produce an assembler error at best and wrong code at worst.",
    "All six relational operators, each yielding exactly 1 or 0.",
    "Order of operations is correct and demonstrable: relational operators bind LOOSER "
    "than arithmetic, so a + 1 < b * 2 groups as (a+1) < (b*2). Your tests must "
    "include an expression that changes value under the wrong precedence.",
    "break leaves the innermost loop. Nested loops must be tested.",
    "break outside any loop is a SEMANTIC error with a line number, not a crash and "
    "not silence.",
    "At least FOUR optimization techniques, each of which the compiler names and "
    "counts. Constant folding, constant propagation, dead code elimination and "
    "algebraic simplification are the natural four.",
    "The optimizer runs to a FIXED POINT — repeatedly until a pass changes nothing — "
    "and reports how many passes it took. A single pass will miss most of what it "
    "could do, and demonstrating that is part of the assignment.",
    "The optimizer forgets everything it knows at a label. Failing to do this produces "
    "a compiler that is correct on straight-line code and silently wrong on loops.",
    "A measured performance comparison: TAC instruction count and SPIM executed-"
    "instruction count, optimized versus unoptimized, on at least one benchmark with a "
    "hot loop. Report both, and explain why they differ.",
    "All Topic 2 and Topic 3 tests still pass.",
  ],
  optional=[
    ("A second loop form",
     "If you implemented while, add for — or the reverse. State whether you lowered it "
     "directly or desugared it into the form you already had, and defend the choice. "
     "Both are legitimate and the reasoning is what is being assessed."),
    ("do-while / repeat-until",
     "The body runs before the first test. Two lines different from while, and it "
     "makes the test-at-top decision concrete."),
    ("continue",
     "Harder than it looks: in a while loop it jumps to the top, but in a for loop it "
     "must jump to the UPDATE or the loop variable never advances. Requires a second "
     "label stack."),
    ("Labelled break",
     "break outer; as in Java. What does your label stack need to become?"),
    ("Loop unrolling",
     "Implement it for loops with a compile-time-known trip count. Measure the effect "
     "on both code size and executed instructions, and report the case where it loses."),
    ("Strength reduction",
     "Replace a multiply in a loop with an add. Your array indexing already does a "
     "simple form of this — find it, then generalise it."),
    ("Loop-invariant code motion",
     "Hoist a computation that does not change out of the loop. The hardest of the "
     "optional optimizations, and the most rewarding to measure."),
  ],
  clarified=[
    ("\"Support for at least one type of loop\"",
     "One is the requirement. Implementing two is optional and should be presented as "
     "a design comparison rather than as extra volume."),
    ("\"Support for order of operations\"",
     "This means the FULL precedence table: unary, then * /, then + -, then relational. "
     "It must be demonstrated with test programs whose output changes if the precedence "
     "is wrong. A test that passes under both correct and incorrect precedence "
     "demonstrates nothing."),
    ("\"Optimize intermediate code\"",
     "Four techniques minimum, each named and counted by the compiler itself. The "
     "compiler must report what it did; a claim in the README that is not backed by "
     "the compiler's own output is not evidence."),
    ("\"Quantify the performance gain from optimization\"",
     "Numbers, with a stated baseline, a stated benchmark and a stated method. Repeat "
     "each measurement at least three times and report the spread — a 5% improvement "
     "with 8% run-to-run variation is not a result. Say whether you are reporting code "
     "SIZE or work DONE; they are different numbers and they will disagree."),
    ("\"Determine language design choices to simplify compilation\"",
     "Write down at least two choices you made and what each one bought you. Examples: "
     "requiring braces on loop bodies, disallowing assignment inside expressions, "
     "having no comma operator. Each removes a class of parsing or lowering problem."),
  ],
  video=[
    ("Describe the language design choices you made",
     "The specific question this project's video asks. Name at least two, say what "
     "each simplified, and say what you gave up."),
    ("Show a loop through the whole pipeline",
     "Source, AST, TAC with labels, optimized TAC, MIPS, and the program running. The "
     "labels are the part worth dwelling on — they are the only reason it is a loop "
     "and not a straight line."),
    ("Present your measurements",
     "Show the table. Say what you compared, on what benchmark, and how many times you "
     "ran it."),
  ],
  start=[
    "Starter code: student/topic-4-loops/compiler/",
    "Lecture notes: docs/topic-4-loops/lecture-notes.html",
    "The TODOs cover scanner tokens, grammar rules, two AST nodes, the break-depth "
    "check, the lowering, the branch code generation, and branch simplification.",
    "Do the measurement work in week 11's activity, not the night before the deadline "
    "— it takes longer than you expect and it is worth marks in Project 6 as well.",
  ]),

 dict(
  n=5, topic=5,
  slug="topic-5-decisions",
  name="Project 5 — Logic and Decisions",
  window="Topic 5 · Nov 23 – Dec 13, 2026",
  due="Saturday, December 13, 2026, 11:59 PM",
  clc=True,
  summary=(
    "Add conditional execution and Boolean operators. The mechanics are the ones you "
    "built for loops; what is new is one famous grammar ambiguity and one real code-"
    "generation choice. This milestone completes the language."),
  goal=(
    "A compiler handling if, if-else and nested if with the dangling-else ambiguity "
    "correctly resolved, plus Boolean operators — and, optionally, switch."),
  new=[
    ("if / if-else", "Lowered with IF_FALSE and GOTO."),
    ("Nested if", "Including the dangling-else case."),
    ("Boolean operators", "&& || ! with C-style truthiness and 0/1 results."),
    ("switch (optional)", "Dispatch chain, fall-through, default, and break."),
    ("Retargeting analysis", "What would change for ARM, and what would not."),
  ],
  required=[
    "if (cond) stmt and if (cond) stmt else stmt, both working.",
    "Nested if statements to at least three levels, tested.",
    "The dangling-else ambiguity RESOLVED EXPLICITLY, not left to bison's default. "
    "Your grammar must produce NO shift/reduce conflicts. Run bison -v and include the "
    "clean output as evidence.",
    "A test program that distinguishes the two readings of if (a) if (b) X else Y, "
    "with the four combinations of a and b, and documented expected output for each.",
    "The three logical operators && || ! , with C-style semantics: any non-zero value "
    "is true, and the result is exactly 0 or 1.",
    "2 && 1 must evaluate to 1. If it evaluates to 0 you emitted a bitwise instruction "
    "without normalising the operands, and every truthy-but-not-1 value is wrong.",
    "Your README states whether your && and || short-circuit, and why you chose that. "
    "Either answer is acceptable; not knowing which you implemented is not.",
    "Semantic analysis warns about a condition that is a compile-time constant — a "
    "constant-false if makes the then-branch unreachable, and the programmer probably "
    "did not mean it.",
    "An audit of your own AST: for at least three non-trivial programs, show the AST "
    "your parser builds and confirm by inspection that it matches the source. Include "
    "these in your submission.",
    "A written comparison of MIPS and ARM code generation for at least one function: "
    "the two listings side by side, the instruction counts, and — most importantly — a "
    "file-by-file statement of what in YOUR compiler would have to change.",
    "All Topic 2, 3 and 4 tests still pass.",
  ],
  optional=[
    ("switch / case / default",
     "Named as optional in the syllabus, and strongly recommended. Requires: a "
     "controlling expression evaluated exactly ONCE, a dispatch chain, correct "
     "fall-through when a case has no break, an optional default, and break jumping "
     "to the end of the switch rather than out of an enclosing loop. Also add the two "
     "semantic checks it makes possible: at most one default, and no duplicate case "
     "values."),
    ("Jump-table dispatch for switch",
     "Instead of a linear comparison chain. Include a written cost comparison for "
     "dense versus sparse case values — this is the more interesting half of the "
     "exercise."),
    ("Short-circuit evaluation",
     "If you implemented eager evaluation, implement short-circuiting as well and "
     "measure the difference on a benchmark where the first operand is usually false."),
    ("The ternary operator",
     "c ? a : b. A small grammar addition with a lowering identical to if-else, and a "
     "good demonstration that you understand the pattern rather than having memorised "
     "one instance of it."),
    ("else if chains as a distinct construct",
     "Most languages treat else if as nesting. Is that what yours does? Show the AST."),
    ("An actual ARM back end",
     "Beyond the written comparison. Even a partial one covering arithmetic and "
     "branches is a substantial and impressive result, and it proves the IR boundary "
     "was real."),
  ],
  clarified=[
    ("\"Support for decision controls: if-then, if-then-else, nested if (optional "
     "switch)\"",
     "if and if-else and nesting are required. switch is optional but recommended — "
     "and if you implement it, fall-through and break are part of it, not extras."),
    ("\"Support for logical (Boolean) operations\"",
     "&& || and ! at minimum. State your truthiness rule and your short-circuit "
     "decision in the README."),
    ("\"Audit the correctness of the AST and semantic actions of the parser\"",
     "A deliberate activity, not a feeling. Show the AST for several non-trivial "
     "programs and confirm it matches the source. The Week 14 activity has you audit "
     "each other's compilers adversarially; use what you find."),
    ("\"Compare code generation for MIPS vs ARM processors\"",
     "A written comparison with concrete instruction sequences for the same function. "
     "The useful part is not the mnemonic table — it is naming which files of your "
     "compiler would change and which would not, and saying what that tells you about "
     "why the compiler has an intermediate representation at all."),
    ("\"Improve and broaden the IR code optimization capabilities\"",
     "At least one optimization beyond the four required in Project 4, or a measurable "
     "improvement to an existing one. Report the before and after."),
  ],
  video=[
    ("Describe how to increase the complexity of language structures a compiler can "
     "handle",
     "The specific question this project's video asks. The strongest answer is "
     "structural: show that adding if touched the same six places that adding while "
     "touched, and that the interfaces between phases did not move."),
    ("Explain the dangling-else problem and your resolution",
     "Show the conflict, show the fix, and show the test program that distinguishes "
     "the two readings."),
    ("Present the MIPS vs ARM comparison",
     "Both listings on screen, and your file-by-file assessment."),
  ],
  start=[
    "Starter code: student/topic-5-decisions/compiler/",
    "Lecture notes: docs/topic-5-decisions/lecture-notes.html",
    "The dangling-else fix is four lines and the order of those four lines is the whole "
    "answer. Build it once WITHOUT them and read bison's conflict report first.",
    "If you attempt switch, do the lowering on paper before writing code — the "
    "dispatch-then-bodies structure is what makes fall-through free.",
  ]),

 dict(
  n=6, topic=6,
  slug="topic-6-complete-compiler",
  name="Project 6 — Complete Compiler (Benchmark)",
  window="Topic 6 · Dec 14 – Dec 20, 2026",
  due="Saturday, December 20, 2026, 11:59 PM",
  clc=False,
  summary=(
    "No new language features. This milestone is about being able to say something "
    "true about the compiler you built: how fast it is, how fast its output is, what "
    "it cannot do, and why it is shaped the way it is. Completed as a group, submitted "
    "individually by each student."),
  goal=(
    "A finished, documented, measured compiler — and a README that lets a stranger "
    "build and run it without asking you anything."),
  new=[
    ("A comprehensive README", "Installation, usage, the grammar, the architecture, "
     "performance, limitations."),
    ("Performance metrics", "Compilation time and execution time, measured and "
     "reported honestly."),
    ("A limitations section", "The honest list, each with a reason."),
    ("An individual submission", "Group work, individual submission and individual "
     "video."),
  ],
  required=[
    "A README that a person who has never seen your compiler can follow to build it "
    "and run it. It must contain: what the compiler does; exact requirements including "
    "TOOL VERSIONS; the exact build command; the exact run command; at least one "
    "complete worked example with its expected output; the grammar in BNF; a "
    "description of each of the six phases naming the file that implements it; the "
    "performance results; and the known limitations.",
    "The README is verified by someone outside your team following it on a clean "
    "checkout. Record who tested it and what they had to guess.",
    "COMPILATION TIME metrics: per-phase timings on inputs of at least three different "
    "sizes. Identify which phase grows fastest and say whether anything is "
    "super-linear.",
    "EXECUTION TIME metrics: SPIM instruction counts for optimized versus unoptimized "
    "output on at least three benchmarks of different shapes (straight-line, hot loop, "
    "call-heavy).",
    "Every measurement repeated at least three times, with the spread reported. State "
    "explicitly what you compared against and on what hardware.",
    "A limitations section listing at least five real limits of your compiler, each "
    "with the reason it exists and what it would take to lift.",
    "A complete regression run: every test from Topics 1 through 5 passes, and the "
    "output is included as evidence.",
    "Error handling demonstrated across all three categories — lexical, syntax and "
    "semantic — with a test file for each and the actual messages shown.",
    "The repository is clean: no build artefacts, no commented-out code, no dead files. "
    "The commit history shows how the work was divided.",
  ],
  optional=[
    ("A --dot flag emitting GraphViz",
     "AST diagrams for your documentation, generated rather than drawn."),
    ("A -O0/-O1/-O2 flag",
     "Optimization levels, so the comparison is a command-line switch rather than a "
     "source edit. Makes your measurements reproducible by someone else, which is the "
     "point of measurements."),
    ("A regression script",
     "One command that compiles every test, runs it in SPIM, and diffs against "
     "expected output. Turns your test suite into something that can catch a "
     "regression rather than merely document one."),
    ("Memory-leak analysis",
     "Run under valgrind and report. The AST and TAC lists are never freed in the "
     "given design; quantify it and say whether it matters for a compiler process that "
     "exits."),
    ("Any Topic 3–5 optional feature not yet attempted",
     "Structures, strings, the ternary operator, short-circuit evaluation, an ARM back "
     "end. Document it as clearly as the required work."),
    ("A one-page language reference",
     "Written for a programmer who wants to USE your language rather than read its "
     "compiler. A genuinely different piece of writing, and a good portfolio artefact."),
  ],
  clarified=[
    ("\"Comprehensive README file with installation and usage instructions\"",
     "The test is operational, not aesthetic: can a stranger build and run it from the "
     "README alone? Have someone outside the team actually try, and record what they "
     "had to guess."),
    ("\"Performance metrics, to include compilation time and execution time\"",
     "Two distinct measurements. Compilation time is how long YOUR COMPILER takes to "
     "run. Execution time is how long ITS OUTPUT takes to run. Both are required, and "
     "each needs a stated method and a stated baseline. Note that SPIM is not a "
     "timing-accurate simulator, so report executed-instruction counts rather than "
     "wall-clock seconds for execution."),
    ("\"Although this project will be completed as a group assignment, it will be "
     "individually submitted\"",
     "Each student submits the repository link and their own video. The code may be "
     "shared; the video and the account of your own contribution must be yours."),
    ("\"Analyze compiler performance\" and \"analyze executable code efficiency\"",
     "Again, two different things: the speed of the compiler, and the quality of the "
     "code it emits. Address both, and be explicit about which numbers belong to "
     "which."),
    ("What to do about features that do not work",
     "Document them honestly in the limitations section with what you know about why. "
     "A known, diagnosed failure demonstrates understanding. A hidden one is found in "
     "thirty seconds by anyone who runs the tests."),
  ],
  video=[
    ("Explain your language design decisions",
     "Not a feature list. Two or three decisions where you chose between real "
     "alternatives — short-circuit or eager, jump table or linear chain, desugar for "
     "or lower it directly, two-pass or one-pass semantic analysis — and why."),
    ("Explain your implementation approach",
     "How the work was divided, what the phase boundaries bought you, and what you "
     "would do differently. This is the reflective content Lab Question 29 asks for; "
     "say it out loud here."),
    ("Present the performance results",
     "The table on screen, with the baseline and method stated."),
    ("State the limitations",
     "At least two, with reasons. Ending on a confident, accurate account of what your "
     "compiler cannot do is a stronger finish than a claim that it does everything."),
  ],
  start=[
    "Starter code: student/topic-6-complete-compiler/compiler/",
    "tests/t6_06_everything.cm exercises every feature at once — if it runs, the "
    "language is complete.",
    "tests/t6_07_benchmark.cm is a starting point for the measurements.",
    "Do the Week 15 activities: the benchmark session produces your performance "
    "section, and the cold-start session finds the holes in your README while there is "
    "still time to fix them.",
  ]),
]


# ===========================================================================
# Rendering
# ===========================================================================

def style(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.12


def heading(doc, text, size=14, color=ACCENT, space_before=16):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return para


def bullets(doc, items, style_name="List Bullet"):
    for it in items:
        para = doc.add_paragraph(style=style_name)
        para.paragraph_format.space_after = Pt(3)
        para.add_run(it)


def pairs(doc, items, numbered=False):
    """A bolded lead-in followed by explanatory text."""
    for i, (lead, body) in enumerate(items, 1):
        para = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        para.paragraph_format.space_after = Pt(4)
        r = para.add_run(lead + " — ")
        r.bold = True
        para.add_run(body)


def build(pr):
    doc = Document()
    style(doc)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)

    # ---- title block ----
    e = doc.add_paragraph()
    e.paragraph_format.space_after = Pt(0)
    r = e.add_run("CST-405 · PRINCIPLES OF COMPILER DESIGN")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(2)
    r = t.add_run(pr["name"])
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = ACCENT

    s = doc.add_paragraph()
    r = s.add_run("%s   ·   Due %s   ·   %s"
                  % (pr["window"], pr["due"],
                     "Collaborative Learning Community (team)" if pr["clc"]
                     else "Group work, submitted individually by each student"))
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED

    doc.add_paragraph()

    # ---- overview ----
    heading(doc, "Overview", space_before=4)
    doc.add_paragraph(pr["summary"])
    para = doc.add_paragraph()
    r = para.add_run("What you are building: ")
    r.bold = True
    para.add_run(pr["goal"])

    # ---- what's new ----
    heading(doc, "What this milestone adds")
    pairs(doc, pr["new"])

    # ---- required ----
    heading(doc, "Required functionality")
    doc.add_paragraph(
        "Every item below must be implemented and demonstrable. Where an item says "
        "\"tested\", a test program for it must be in your repository with its "
        "expected result stated in a header comment.")
    bullets(doc, pr["required"])

    # ---- optional ----
    heading(doc, "Optional extensions")
    doc.add_paragraph(
        "None of these is required, and none of them substitutes for anything in the "
        "required list. Attempt them only after the required work is complete and "
        "tested. If you claim an optional feature, it must be documented in the README "
        "and demonstrated in your video to the same standard as the required work — an "
        "undocumented half-feature is worth less than not attempting it.")
    pairs(doc, pr["optional"])

    # ---- clarifications ----
    heading(doc, "Clarified expectations")
    doc.add_paragraph(
        "These points expand on wording in the course syllabus that has more than one "
        "reasonable reading. Where this document and the syllabus differ in detail, "
        "this document is what the work is assessed against.")
    pairs(doc, pr["clarified"])

    # ---- code documentation ----
    heading(doc, "Code documentation requirements")
    doc.add_paragraph(
        "Documented code is a requirement of this assignment, not a courtesy. The "
        "standard is that a reader who knows compilers but has never seen your code "
        "can follow the design without asking you anything.")
    bullets(doc, DOC_REQUIREMENTS)

    # ---- github ----
    heading(doc, "Repository")
    bullets(doc, GITHUB)

    # ---- video ----
    heading(doc, "Video requirement — one per team member")
    doc.add_paragraph(VIDEO_INTRO)
    p2 = doc.add_paragraph()
    r = p2.add_run("Every video must cover the following:")
    r.bold = True
    pairs(doc, VIDEO_COMMON, numbered=True)
    p3 = doc.add_paragraph()
    r = p3.add_run("And, specific to this project:")
    r.bold = True
    pairs(doc, pr["video"], numbered=True)
    doc.add_paragraph(
        "Suggested length: 8–12 minutes. Screen-share the code and the running "
        "compiler throughout. Reading code aloud is the most common way these videos "
        "go wrong — the viewer can read. Spend the time on why it is shaped that way, "
        "what you tried that did not work, and what you would do differently.")

    # ---- submission ----
    heading(doc, "Submission checklist")
    bullets(doc, [
        "Link to the private GitHub repository, with the instructor added as a "
        "collaborator.",
        "One video link per team member.",
        "A short document listing which team member wrote which phase.",
        "Any screenshots the required-functionality list asks for, each with a caption "
        "explaining what it shows.",
        "Confirmation that the previous topics' tests still pass (from Project 3 "
        "onward).",
    ])
    doc.add_paragraph(
        "APA style is not required. Solid academic writing is expected. You are not "
        "required to submit this assignment to LopesWrite.")

    # ---- where to start ----
    heading(doc, "Where to start")
    bullets(doc, pr["start"])

    # ---- footer note ----
    doc.add_paragraph()
    f = doc.add_paragraph()
    r = f.add_run(
        "This assignment reinforces the competency: apply computer science theory and "
        "software development fundamentals to produce computing-based solutions."
        if pr["n"] != 6 else
        "Benchmark assignment. Assesses programmatic competencies 4.1, 4.2, 4.3 "
        "(ABET 3), 6.1 and 7.4 (ABET 5).")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED

    out = os.path.join(DOCS, pr["slug"],
                       "CST-405-Topic-%d-Project-%d.docx" % (pr["topic"], pr["n"]))
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    return out


if __name__ == "__main__":
    for pr in PROJECTS:
        print("wrote", os.path.relpath(build(pr), ROOT))
